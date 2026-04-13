from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

TEMPLATE_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
)
DOCUMENT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
)
PACKAGE_PLACEHOLDER_DEFAULTS = {
    "{{ENGAGEMENT_TITLE}}": "ENGAGEMENT TITLE",
    "{{PROPOSAL_NUMBER}}": "PROPOSAL NUMBER",
    "{{CLIENT}}": "CLIENT",
    "{{DUE_DATE}}": "DUE DATE",
    "{{DUE_TIME}}": "DUE TIME",
}
RELATIONSHIP_NAMESPACE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
)
WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
RELATIONSHIP_ATTRIBUTE_NAMES = {
    f"{{{RELATIONSHIP_NAMESPACE}}}embed",
    f"{{{RELATIONSHIP_NAMESPACE}}}id",
    f"{{{RELATIONSHIP_NAMESPACE}}}link",
}
PROJECT_BORDER_START_COLOR = "494949"
PROJECT_BORDER_END_COLOR_SHORT_LIST = "BDBDBD"
PROJECT_BORDER_END_COLOR_LONG_LIST = "D9D9D9"
PROJECT_BORDER_THRESHOLD_FOR_LONG_LIST = 5


def _convert_template_to_docx(
    template_path: str | Path, destination_path: str | Path
) -> None:
    with ZipFile(str(template_path), "r") as source_zip, ZipFile(
        str(destination_path), "w", ZIP_DEFLATED
    ) as destination_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    TEMPLATE_CONTENT_TYPE.encode("utf-8"),
                    DOCUMENT_CONTENT_TYPE.encode("utf-8"),
                )
            destination_zip.writestr(item, data)


def _iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_placeholder_in_runs(
    paragraph: Paragraph, placeholder: str, value: str
) -> bool:
    if not paragraph.runs:
        return False

    full_text = "".join(run.text for run in paragraph.runs)
    start = full_text.find(placeholder)
    if start == -1:
        return False

    end = start + len(placeholder)
    char_index = 0
    start_run_index = None
    end_run_index = None

    for run_index, run in enumerate(paragraph.runs):
        next_char_index = char_index + len(run.text)
        if start_run_index is None and start < next_char_index:
            start_run_index = run_index
        if end <= next_char_index:
            end_run_index = run_index
            break
        char_index = next_char_index

    if start_run_index is None or end_run_index is None:
        return False

    run_start_offsets: list[int] = []
    char_index = 0
    for run in paragraph.runs:
        run_start_offsets.append(char_index)
        char_index += len(run.text)

    for run_index in range(start_run_index, end_run_index + 1):
        run = paragraph.runs[run_index]
        run_text = run.text
        run_global_start = run_start_offsets[run_index]
        run_global_end = run_global_start + len(run_text)
        local_start = max(start, run_global_start) - run_global_start
        local_end = min(end, run_global_end) - run_global_start
        replacement_targets = [
            idx
            for idx in range(start_run_index, end_run_index + 1)
            if paragraph.runs[idx].text.strip("{}")
        ]
        target_run_index = (
            max(replacement_targets, key=lambda idx: len(paragraph.runs[idx].text))
            if replacement_targets
            else start_run_index
        )

        if run_index == target_run_index:
            run.text = run_text[:local_start] + value + run_text[local_end:]
        else:
            run.text = run_text[:local_start] + run_text[local_end:]

    return True


def _replace_placeholders_in_paragraph(
    paragraph: Paragraph, replacements: dict[str, str]
) -> None:
    for placeholder, value in replacements.items():
        while _replace_placeholder_in_runs(paragraph, placeholder, value):
            pass


def _replace_placeholders_in_table(table: Table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_in_paragraph(paragraph, replacements)


def _remove_spacing_for_paragraphs_in_table(table: Table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)


def _hex_to_rgb(color: str) -> tuple[int, int, int]:
    normalized = color.strip().lstrip("#").upper()
    if len(normalized) != 6:
        raise ValueError(f"Expected a 6-character hex color, got '{color}'.")
    return tuple(int(normalized[index : index + 2], 16) for index in range(0, 6, 2))


def _rgb_to_hex(color: tuple[int, int, int]) -> str:
    return "".join(f"{channel:02X}" for channel in color)


def _interpolate_hex_color(start: str, end: str, progress: float) -> str:
    start_rgb = _hex_to_rgb(start)
    end_rgb = _hex_to_rgb(end)
    channels = tuple(
        round(start_channel + (end_channel - start_channel) * progress)
        for start_channel, end_channel in zip(start_rgb, end_rgb)
    )
    return _rgb_to_hex(channels)


def _build_project_border_colors(project_count: int) -> list[str]:
    if project_count <= 0:
        return []
    if project_count == 1:
        return [PROJECT_BORDER_START_COLOR]

    end_color = (
        PROJECT_BORDER_END_COLOR_LONG_LIST
        if project_count >= PROJECT_BORDER_THRESHOLD_FOR_LONG_LIST
        else PROJECT_BORDER_END_COLOR_SHORT_LIST
    )
    return [
        _interpolate_hex_color(
            PROJECT_BORDER_START_COLOR,
            end_color,
            index / (project_count - 1),
        )
        for index in range(project_count)
    ]


def _get_or_add_child(parent, tag_name: str):
    child = parent.find(qn(tag_name))
    if child is None:
        child = OxmlElement(tag_name)
        parent.append(child)
    return child


def _set_border_element_color(border_element, color: str) -> None:
    border_element.set(qn("w:color"), color)
    border_element.attrib.pop(qn("w:themeColor"), None)
    border_element.attrib.pop(qn("w:themeShade"), None)


def _apply_right_border_color_to_project_table(table: Table, color: str) -> None:
    normalized_color = color.strip().lstrip("#").upper()
    if len(table.rows) < 1 or len(table.rows[0].cells) < 2:
        return

    title_cell = table.rows[0].cells[0]
    description_cell = table.rows[0].cells[1]

    title_tc_pr = _get_or_add_child(title_cell._tc, "w:tcPr")
    title_tc_borders = _get_or_add_child(title_tc_pr, "w:tcBorders")
    right_border = _get_or_add_child(title_tc_borders, "w:right")
    _set_border_element_color(right_border, normalized_color)

    description_tc_pr = _get_or_add_child(description_cell._tc, "w:tcPr")
    description_tc_borders = _get_or_add_child(description_tc_pr, "w:tcBorders")
    left_border = _get_or_add_child(description_tc_borders, "w:left")
    _set_border_element_color(left_border, normalized_color)


def _replace_placeholders_in_document(
    document: Document, replacements: dict[str, str]
) -> None:
    for paragraph in document.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, replacements)
    for table in document.tables:
        _replace_placeholders_in_table(table, replacements)


def _iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _format_person_title(document: Document, title: str) -> None:
    normalized_title = title.strip()
    if not normalized_title:
        return

    uppercase_title = normalized_title.upper()
    for paragraph in _iter_all_paragraphs(document):
        if paragraph.text.strip() != normalized_title:
            continue

        _set_paragraph_text(paragraph, uppercase_title)
        for run in paragraph.runs:
            run.font.size = Pt(14)


def _format_person_name(document: Document, first_name: str, last_name: str) -> None:
    full_name = f"{first_name.strip()} {last_name.strip()}".strip()
    if not full_name:
        return

    for paragraph in _iter_all_paragraphs(document):
        if paragraph.text.strip() != full_name:
            continue
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)


def _find_paragraph_index(paragraphs: list[Paragraph], marker: str) -> int:
    for index, paragraph in enumerate(paragraphs):
        if marker in paragraph.text:
            return index
    raise ValueError(f"Marker '{marker}' was not found in the template.")


def _find_block_range(
    document: Document,
    start_marker: str,
    end_marker: str,
) -> tuple[list[Paragraph | Table], int, int]:
    block_items = list(_iter_block_items(document))
    start_index = next(
        (
            index
            for index, item in enumerate(block_items)
            if isinstance(item, Paragraph) and item.text == start_marker
        ),
        None,
    )
    end_index = next(
        (
            index
            for index, item in enumerate(block_items)
            if isinstance(item, Paragraph) and item.text == end_marker
        ),
        None,
    )

    if start_index is None or end_index is None or end_index <= start_index:
        raise ValueError(
            f"Unable to locate block markers {start_marker} and {end_marker}."
        )

    return block_items, start_index, end_index


def _try_find_block_range(
    document: Document,
    start_marker: str,
    end_marker: str,
) -> tuple[list[Paragraph | Table], int, int] | None:
    try:
        return _find_block_range(document, start_marker, end_marker)
    except ValueError:
        return None


def _remove_items(items: list[Paragraph | Table]) -> None:
    for item in items:
        item._element.getparent().remove(item._element)


def _render_optional_block(
    document: Document,
    start_marker: str,
    end_marker: str,
    include_block: bool,
) -> None:
    block_range = _try_find_block_range(document, start_marker, end_marker)
    if block_range is None:
        if include_block:
            raise ValueError(
                f"Unable to locate block markers {start_marker} and {end_marker}."
            )
        return

    block_items, start_index, end_index = block_range
    if include_block:
        _remove_items([block_items[start_index], block_items[end_index]])
        return
    _remove_items(block_items[start_index : end_index + 1])


def _collect_block_elements(
    document: Document,
    start_marker: str,
    end_marker: str,
) -> list:
    block_items, start_index, end_index = _find_block_range(
        document, start_marker, end_marker
    )
    return [
        deepcopy(item._element) for item in block_items[start_index + 1 : end_index]
    ]


def _build_package_replacements(package_context: dict | None) -> dict[str, str]:
    context = package_context or {}
    field_map = {
        "{{ENGAGEMENT_TITLE}}": "engagement_title",
        "{{PROPOSAL_NUMBER}}": "proposal_number",
        "{{CLIENT}}": "client",
        "{{DUE_DATE}}": "due_date",
        "{{DUE_TIME}}": "due_time",
    }
    replacements: dict[str, str] = {}
    for placeholder, field_name in field_map.items():
        value = context.get(field_name)
        replacements[placeholder] = (
            str(value).strip()
            if value
            else PACKAGE_PLACEHOLDER_DEFAULTS.get(placeholder, placeholder)
        )
    return replacements


def _render_repeating_block(
    document: Document,
    start_marker: str,
    end_marker: str,
    items: list[dict[str, str]],
    placeholder_map: dict[str, str],
) -> None:
    block_items, start_index, end_index = _find_block_range(
        document, start_marker, end_marker
    )
    template_items = block_items[start_index + 1 : end_index]
    insert_after = block_items[start_index]._element
    project_border_colors = (
        _build_project_border_colors(len(items))
        if start_marker == "{{PROJECT_BLOCK_START}}"
        else []
    )

    for index, item in enumerate(items):
        replacements = {
            placeholder: str(item.get(field_name, ""))
            for placeholder, field_name in placeholder_map.items()
        }
        for template_item in template_items:
            new_element = deepcopy(template_item._element)
            insert_after.addnext(new_element)
            insert_after = new_element
            wrapper = (
                Paragraph(new_element, document)
                if new_element.tag.endswith("}p")
                else Table(new_element, document)
            )
            if isinstance(wrapper, Paragraph):
                _replace_placeholders_in_paragraph(wrapper, replacements)
            else:
                _replace_placeholders_in_table(wrapper, replacements)
                if start_marker == "{{PROJECT_BLOCK_START}}":
                    _remove_spacing_for_paragraphs_in_table(wrapper)
                    _apply_right_border_color_to_project_table(
                        wrapper, project_border_colors[index]
                    )

    for template_item in template_items:
        template_item._element.getparent().remove(template_item._element)
    _remove_items([block_items[start_index], block_items[end_index]])


def _load_document(path: str | Path) -> Document:
    source_path = Path(path)
    if source_path.suffix.lower() == ".dotx":
        with NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_path = Path(temp_file.name)
        _convert_template_to_docx(source_path, temp_path)
        try:
            return Document(str(temp_path))
        finally:
            temp_path.unlink(missing_ok=True)
    return Document(str(source_path))


def _copy_relationship_to_document(
    source_document: Document,
    target_document: Document,
    source_rel_id: str,
    rel_id_map: dict[str, str],
) -> str:
    if source_rel_id in rel_id_map:
        return rel_id_map[source_rel_id]

    relationship = source_document.part.rels[source_rel_id]
    if relationship.is_external:
        target_rel_id = target_document.part.relate_to(
            relationship.target_ref,
            relationship.reltype,
            is_external=True,
        )
    elif relationship.reltype == RT.IMAGE:
        target_rel_id, _ = target_document.part.get_or_add_image(
            BytesIO(relationship.target_part.blob)
        )
    else:
        raise ValueError(
            f"Unsupported internal relationship type '{relationship.reltype}' while merging document content."
        )

    rel_id_map[source_rel_id] = target_rel_id
    return target_rel_id


def _clone_element_for_merge(
    source_document: Document,
    target_document: Document,
    element,
    rel_id_map: dict[str, str],
):
    cloned = deepcopy(element)
    for node in cloned.iter():
        for attr_name, attr_value in list(node.attrib.items()):
            if (
                attr_name not in RELATIONSHIP_ATTRIBUTE_NAMES
                or attr_value not in source_document.part.rels
            ):
                continue
            node.set(
                attr_name,
                _copy_relationship_to_document(
                    source_document, target_document, attr_value, rel_id_map
                ),
            )
    return cloned


def _prepend_elements_to_document(
    document: Document,
    elements: list,
    source_document: Document | None = None,
) -> None:
    if not elements:
        return

    body = document.element.body
    first_child = next(
        (child for child in body.iterchildren() if not child.tag.endswith("}sectPr")),
        None,
    )
    rel_id_map: dict[str, str] = {}
    if first_child is None:
        for element in elements:
            element_to_add = (
                _clone_element_for_merge(source_document, document, element, rel_id_map)
                if source_document is not None
                else deepcopy(element)
            )
            body.append(element_to_add)
        return

    anchor = first_child
    for element in elements:
        element_to_add = (
            _clone_element_for_merge(source_document, document, element, rel_id_map)
            if source_document is not None
            else deepcopy(element)
        )
        anchor.addprevious(element_to_add)


def _append_elements_to_document(
    document: Document,
    elements: list,
    source_document: Document | None = None,
) -> None:
    if not elements:
        return

    body = document.element.body
    section_properties = next(
        (child for child in body.iterchildren() if child.tag.endswith("}sectPr")), None
    )
    rel_id_map: dict[str, str] = {}
    if section_properties is None:
        for element in elements:
            element_to_add = (
                _clone_element_for_merge(source_document, document, element, rel_id_map)
                if source_document is not None
                else deepcopy(element)
            )
            body.append(element_to_add)
        return

    for element in elements:
        element_to_add = (
            _clone_element_for_merge(source_document, document, element, rel_id_map)
            if source_document is not None
            else deepcopy(element)
        )
        section_properties.addprevious(element_to_add)


def _collect_document_body_elements(document: Document) -> list:
    return [
        deepcopy(child)
        for child in document.element.body.iterchildren()
        if not child.tag.endswith("}sectPr")
    ]


def _collect_resume_body_elements(document: Document) -> list:
    start_marker = "{{RESUME_START}}"
    end_marker = "{{RESUME_END}}"
    block_range = _try_find_block_range(document, start_marker, end_marker)
    if block_range is not None:
        block_items, start_index, end_index = block_range
        elements = [
            deepcopy(item._element) for item in block_items[start_index + 1 : end_index]
        ]
        return _strip_boundary_page_break_paragraphs(elements)

    block_items = list(_iter_block_items(document))
    has_start = any(
        isinstance(item, Paragraph) and item.text == start_marker
        for item in block_items
    )
    has_end = any(
        isinstance(item, Paragraph) and item.text == end_marker for item in block_items
    )
    if has_start or has_end:
        raise ValueError(
            "Resume boundary markers are incomplete. Ensure both '{{RESUME_START}}' and "
            "'{{RESUME_END}}' exist as standalone paragraphs in the template."
        )

    # Backward compatibility for templates that do not yet include explicit resume boundaries.
    return _strip_boundary_page_break_paragraphs(
        _collect_document_body_elements(document)
    )


def _is_page_break_only_paragraph(element) -> bool:
    if not element.tag.endswith("}p"):
        return False

    has_page_break = False
    for node in element.iter():
        if node.tag.endswith("}t") and (node.text or "").strip():
            return False
        if node.tag.endswith("}br"):
            break_type = node.get(f"{{{WORD_NAMESPACE}}}type")
            if break_type == "page":
                has_page_break = True
            else:
                return False

    return has_page_break


def _element_contains_page_break(element) -> bool:
    for node in element.iter():
        if node.tag.endswith("}br") and node.get(f"{{{WORD_NAMESPACE}}}type") == "page":
            return True
    return False


def _chunk_has_leading_page_break(elements: list) -> bool:
    if not elements:
        return False
    return _element_contains_page_break(elements[0])


def _chunk_has_trailing_page_break(elements: list) -> bool:
    if not elements:
        return False
    return _element_contains_page_break(elements[-1])


def _create_page_break_element(document: Document):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    element = deepcopy(paragraph._element)
    paragraph._element.getparent().remove(paragraph._element)
    return element


def _strip_chunk_leading_page_breaks(elements: list) -> None:
    if not elements:
        return

    first_element = elements[0]
    page_break_nodes = [
        node
        for node in first_element.iter()
        if node.tag.endswith("}br") and node.get(f"{{{WORD_NAMESPACE}}}type") == "page"
    ]
    for node in page_break_nodes:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def _strip_boundary_page_break_paragraphs(elements: list) -> list:
    start_index = 0
    end_index = len(elements)

    while start_index < end_index and _is_page_break_only_paragraph(
        elements[start_index]
    ):
        start_index += 1
    while end_index > start_index and _is_page_break_only_paragraph(
        elements[end_index - 1]
    ):
        end_index -= 1

    return elements[start_index:end_index]


def _replace_document_body_elements(document: Document, elements: list) -> None:
    body = document.element.body
    existing_content = [
        child for child in body.iterchildren() if not child.tag.endswith("}sectPr")
    ]
    for child in existing_content:
        body.remove(child)
    _append_elements_to_document(document, elements)


def _append_single_page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _load_optional_document_elements(
    path: str | Path,
    replacements: dict[str, str],
) -> tuple[Document, list]:
    document = _load_document(path)
    _replace_placeholders_in_document(document, replacements)
    return document, _collect_document_body_elements(document)


def _render_individual_template(
    template_path: str | Path,
    output_path: str | Path,
    person_data: dict,
    projects: list[dict],
    package_context: dict | None = None,
    include_cover: bool = False,
    include_end_page: bool = False,
) -> None:
    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    _convert_template_to_docx(template_path, output_file)
    document = Document(str(output_file))

    person_replacements = {
        "{{PERSON_FIRST}}": person_data.get("first_name", ""),
        "{{PERSON_LAST}}": person_data.get("last_name", ""),
        "{{PERSON_TITLE}}": person_data.get("title", ""),
        "{{PERSON_SUMMARY}}": person_data.get("summary", ""),
    }

    replacements = {
        **_build_package_replacements(package_context),
        **person_replacements,
    }
    _replace_placeholders_in_document(document, replacements)

    _render_optional_block(
        document=document,
        start_marker="{{COVER_BLOCK_START}}",
        end_marker="{{COVER_BLOCK_END}}",
        include_block=include_cover,
    )
    _render_optional_block(
        document=document,
        start_marker="{{END_BLOCK_START}}",
        end_marker="{{END_BLOCK_END}}",
        include_block=include_end_page,
    )
    _render_optional_block(
        document=document,
        start_marker="{{RESUME_START}}",
        end_marker="{{RESUME_END}}",
        include_block=True,
    )

    _render_repeating_block(
        document=document,
        start_marker="{{PROJECT_BLOCK_START}}",
        end_marker="{{PROJECT_BLOCK_END}}",
        items=projects,
        placeholder_map={
            "{{PROJECT_CLIENT}}": "client",
            "{{PROJECT_TITLE}}": "title",
            "{{PROJECT_DESCRIPTION}}": "description",
        },
    )

    education = person_data.get("education", [])
    _render_repeating_block(
        document=document,
        start_marker="{{EDU_BLOCK_START}}",
        end_marker="{{EDU_BLOCK_END}}",
        items=education,
        placeholder_map={
            "{{DEGREE_CERT}}": "degree_cert",
            "{{DEGREE_AREA}}": "degree_area",
            "{{LOCATION}}": "location",
        },
    )

    _format_person_name(
        document,
        person_data.get("first_name", ""),
        person_data.get("last_name", ""),
    )
    _format_person_title(document, person_data.get("title", ""))

    document.save(str(output_file))


def write_individual_resume(
    person_data: dict,
    projects: list[dict],
    template_path: str | Path,
    output_path: str | Path,
    package_context: dict | None = None,
    include_cover: bool = False,
    include_end_page: bool = False,
) -> None:
    """Create a single-person resume document from the proposal template."""
    _render_individual_template(
        template_path,
        output_path,
        person_data,
        projects,
        package_context=package_context,
        include_cover=include_cover,
        include_end_page=include_end_page,
    )


def write_consolidated_resume(
    individual_paths: list[str | Path],
    template_path: str | Path,
    output_path: str | Path,
    package_context: dict | None = None,
    include_cover: bool = False,
    include_end_page: bool = False,
    cover_page_path: str | Path | None = None,
    end_page_path: str | Path | None = None,
) -> None:
    """Create a consolidated document by stitching individual resume files together."""
    source_paths = [Path(path) for path in individual_paths]
    if not source_paths:
        raise ValueError(
            "At least one individual resume is required to build a consolidated document."
        )

    merged_document = Document(str(source_paths[0]))
    first_elements = _collect_resume_body_elements(merged_document)
    _strip_chunk_leading_page_breaks(first_elements)
    _replace_document_body_elements(merged_document, first_elements)
    cover_document: Document | None = None
    cover_elements: list = []
    end_document: Document | None = None
    end_elements: list = []

    replacements = _build_package_replacements(package_context)
    if include_cover and cover_page_path:
        cover_document, cover_elements = _load_optional_document_elements(
            cover_page_path, replacements
        )
        if cover_elements and not _chunk_has_trailing_page_break(cover_elements):
            cover_elements = [
                *cover_elements,
                _create_page_break_element(merged_document),
            ]
    if include_end_page and end_page_path:
        end_document, end_elements = _load_optional_document_elements(
            end_page_path, replacements
        )

    for source_path in source_paths[1:]:
        source_document = Document(str(source_path))
        source_elements = _collect_resume_body_elements(source_document)
        if not source_elements:
            continue
        _strip_chunk_leading_page_breaks(source_elements)
        _append_single_page_break(merged_document)
        _append_elements_to_document(
            merged_document, source_elements, source_document=source_document
        )

    _prepend_elements_to_document(
        merged_document, cover_elements, source_document=cover_document
    )
    if end_elements and not _chunk_has_leading_page_break(end_elements):
        _append_single_page_break(merged_document)
    _append_elements_to_document(
        merged_document, end_elements, source_document=end_document
    )

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    merged_document.save(str(output_path))
